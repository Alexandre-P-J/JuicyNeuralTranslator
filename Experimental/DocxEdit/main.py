import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.styles.style import _CharacterStyle


document = Document('/home/pyro/Downloads/test.docx')
document = Document('/home/pyro/Downloads/informe-final.docx')

def unfold(doc):
    texts = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                texts.append(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(unfold(cell))
    return texts

def fold(doc, replacements):
    def substitution(ref, replacements, offset):
        for paragraph in ref.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = replacements[offset]
                    offset += 1
        for table in ref.tables:
            for row in table.rows:
                for cell in row.cells:
                    _, offset = substitution(cell, replacements, offset)
        return ref, offset
    return substitution(doc, replacements, 0)[0]


data = unfold(document)
data = [">"+d+"<" for d in data]
document = fold(document, data)
document.save('test.docx')